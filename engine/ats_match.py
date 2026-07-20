#!/usr/bin/env python3
"""
ats_match.py — resume-vs-JD keyword coverage checker.
================================================================================
Scores how well a tailored resume covers the important terms in a specific job
description, flags what's missing, and tells you whether you clear a threshold
(default 85). Run it on a resume before you submit.

RUN:
  python ats_match.py "applications/farmers/Firstname Lastname - Resume.docx" jd.txt
  python ats_match.py "applications/farmers/Firstname Lastname - Resume.docx" --jd-paste
  python ats_match.py "<resume>.docx" jd.txt --min 85

WHAT IT IS (and what it is NOT):
  This is a KEYWORD-COVERAGE heuristic, not a real ATS score. Real ATS platforms
  (Workday, Greenhouse, iCIMS) parse and rank differently, and few publish a
  literal "match %". Treat 85 as a useful bar, not gospel.

  THE HONESTY RULE STILL APPLIES. The right way to raise the score is to surface a
  true term the JD wants that your resume left out - a real skill, in your profile,
  not on the page yet - and add it. It is NOT to stuff keywords you can't back up.
  This tool sorts the gaps into "you can likely claim this" vs "verify - only add
  if true", and never invents anything for you.
"""

import argparse
import os
import re
import sys

import _paths  # noqa: F401  (side-effect: adds repo root to sys.path for `import config`)

# Terms that are high-value when an ATS/recruiter scans a data-leadership resume.
# (Edit for your field. This just tells the scorer which JD words matter most.)
SKILL_LEXICON = {
    "sql",
    "python",
    "dbt",
    "snowflake",
    "databricks",
    "bigquery",
    "redshift",
    "airflow",
    "fivetran",
    "tableau",
    "power bi",
    "sigma",
    "looker",
    "thoughtspot",
    "aws",
    "gcp",
    "azure",
    "vertex ai",
    "sagemaker",
    "medallion",
    "ci/cd",
    "data governance",
    "governance",
    "data quality",
    "metadata",
    "lineage",
    "master data",
    "mdm",
    "reference data",
    "data contracts",
    "stewardship",
    "data platform",
    "data strategy",
    "data management",
    "data engineering",
    "analytics engineering",
    "data mesh",
    "dimensional modeling",
    "data modeling",
    "self-service",
    "enablement",
    "adoption",
    "pii",
    "rbac",
    "compliance",
    "ai",
    "genai",
    "llm",
    "mlops",
    "ml",
    "machine learning",
    "rag",
    "vector",
    "agents",
    "mcp",
    "copilot",
    "cursor",
    "finops",
    "roadmap",
    "stakeholder",
    "leadership",
    "strategy",
    "architecture",
    "reliability",
    "sla",
    "kpi",
    "insurance",
    "financial services",
    "healthcare",
    "saas",
}

STOP = set(
    """a an the and or but if then else for to of in on at by with from as is are was were be been
being this that these those it its it's you your yours we our us they them their he she his her
will would can could should may might must not no yes do does did done have has had having i me my
role position job company team work working experience years year plus etc via across into over under
about within without per your our their more most least very much many few new using used use uses
who whom which what when where why how all any some each other another such same than too also just
including include includes required preferred responsibilities qualifications description ability able
strong excellent proven demonstrated track record help support drive lead build create ensure develop
manage partner deliver best world class high quality""".split()
)


def resume_text(path):
    if path.lower().endswith(".docx"):
        try:
            from docx import Document
        except ImportError:
            sys.exit("Need python-docx to read .docx:  pip install python-docx")
        d = Document(path)
        parts = [p.text for p in d.paragraphs]
        for tbl in d.tables:
            for row in tbl.rows:
                for c in row.cells:
                    parts.append(c.text)
        return "\n".join(parts)
    return open(path, encoding="utf-8", errors="ignore").read()


def norm(s):
    return re.sub(r"[^a-z0-9+/&. ]", " ", s.lower())


def extract_jd_terms(jd):
    """Return {term: weight}. Weight is higher for lexicon skills and repeats,
    and higher again for anything in a requirements/qualifications section."""
    low = jd.lower()
    # find the requirements region (weight terms there 2x)
    req_zone = ""
    m = re.search(
        r"(requirements|qualifications|what you.?ll need|must[- ]have|"
        r"skills|experience required|about you)([\s\S]{0,2500})",
        low,
    )
    if m:
        req_zone = m.group(2)

    n = norm(jd)
    words = n.split()
    terms: dict[str, int] = {}

    # 1) multi-word skills from the lexicon (phrase match)
    for term in SKILL_LEXICON:
        if term in n:
            w = 3 if " " in term else 2  # multiword skills are more specific
            if term in norm(req_zone):
                w += 2
            terms[term] = max(terms.get(term, 0), w)

    # 2) single tokens: keep meaningful ones (frequency or lexicon)
    freq: dict[str, int] = {}
    for tok in words:
        if len(tok) < 3 or tok in STOP or tok.isdigit():
            continue
        freq[tok] = freq.get(tok, 0) + 1
    for tok, f in freq.items():
        if tok in terms:
            continue
        if tok in SKILL_LEXICON or f >= 2:
            w = 2 if tok in SKILL_LEXICON else 1
            if tok in norm(req_zone):
                w += 1
            terms[tok] = w
    return terms


def term_present(term, rtext):
    # flexible: spaces match spaces/hyphens; word-ish boundaries
    pat = re.escape(term).replace(r"\ ", r"[\s\-/]+")
    return re.search(r"(?<![a-z0-9])" + pat + r"(?![a-z0-9])", rtext, re.I) is not None


def evaluate(resume_path, jd, min_score=85):
    """Pure: compute keyword coverage of a resume against a JD. No printing.

    Returns {pct, terms, present, missing, passed, empty} — `present`/`missing`
    are [(term, weight)] sorted by weight desc. `empty` = the JD yielded no terms.
    """
    rtext = norm(resume_text(resume_path))
    terms = extract_jd_terms(jd)
    if not terms:
        return {"pct": 0, "terms": {}, "present": [], "missing": [], "passed": False, "empty": True}

    total_w = sum(terms.values())
    hit_w = 0
    present, missing = [], []
    for term, w in sorted(terms.items(), key=lambda x: -x[1]):
        if term_present(term, rtext):
            hit_w += w
            present.append((term, w))
        else:
            missing.append((term, w))

    pct = round(100 * hit_w / total_w)
    return {
        "pct": pct,
        "terms": terms,
        "present": present,
        "missing": missing,
        "passed": pct >= min_score,
        "empty": False,
    }


def report(result, min_score=85):
    """Print the human-readable coverage report from an evaluate() result."""
    if result["empty"]:
        print("Could not extract terms from the JD (is it empty?).")
        return
    pct, terms, present, missing = (
        result["pct"],
        result["terms"],
        result["present"],
        result["missing"],
    )

    print(f"\n{'=' * 64}")
    print(
        f"  ATS keyword coverage: {pct}%   (threshold {min_score})   "
        f"{'PASS ✅' if result['passed'] else 'BELOW ❌'}"
    )
    print(f"  {len(present)}/{len(terms)} weighted terms matched")
    print(f"{'=' * 64}")

    if missing:
        print("\n  MISSING - highest value first. Add ONLY the ones that are TRUE for you:")
        for t, w in missing[:20]:
            tag = "skill" if t in SKILL_LEXICON else "term"
            print(f"    [{w}] {t}   ({tag})")
        print("\n  Honesty check: a term you can't back up does not go on the resume to")
        print("  clear a threshold. If it's real and traceable to career-profile.md, add it;")
        print("  if it isn't, the honest move is to let the score sit lower.")

    if pct < min_score:
        gap_terms = [t for t, _ in missing]
        print(f"\n  To clear {min_score}%: cover the true items above. The biggest single")
        print("  levers are the multi-word skills you already have but didn't name on THIS")
        print(f"  resume (e.g. {', '.join(gap_terms[:4])}).")


def score(resume_path, jd, min_score=85):
    """Compute + print + return the pct (CLI convenience wrapper)."""
    result = evaluate(resume_path, jd, min_score)
    report(result, min_score)
    return result["pct"]


def main():
    # Optional per-user lexicon override, read only when run as a tool — so the
    # module stays importable (and testable) without a configured personal/.
    try:
        import config

        extra = getattr(config, "ATS_SKILL_LEXICON", None)
        if extra:
            SKILL_LEXICON.update(extra)
    except Exception:  # noqa: S110
        pass

    ap = argparse.ArgumentParser(description="resume-vs-JD keyword coverage check")
    ap.add_argument("resume", help="path to the resume (.docx or .txt)")
    ap.add_argument("jd", nargs="?", help="path to a text file with the job description")
    ap.add_argument(
        "--jd-paste", action="store_true", help="paste the JD on stdin instead of a file"
    )
    ap.add_argument("--min", type=int, default=85, help="pass threshold (default 85)")
    args = ap.parse_args()

    if args.jd_paste or not args.jd:
        print("Paste the job description, then Ctrl-D (Ctrl-Z on Windows):", file=sys.stderr)
        jd = sys.stdin.read()
    else:
        jd = open(args.jd, encoding="utf-8", errors="ignore").read()

    if not os.path.exists(args.resume):
        sys.exit(f"Resume not found: {args.resume}")
    pct = score(args.resume, jd, args.min)
    raise SystemExit(0 if pct >= args.min else 1)


if __name__ == "__main__":
    main()
