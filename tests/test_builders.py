"""Tests for the doc-builder pure logic (F2/F5) — placeholder scanning and the
resume ATS-rule validator. Config-free (needs python-docx only). Stdlib unittest.
"""

import os
import shutil
import sys
import tempfile
import unittest

sys.path.insert(
    0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "engine")
)
import docx_common
import resume_builder


class TestScanPlaceholders(unittest.TestCase):
    def test_finds_stubs_in_nested_content(self):
        spec = {"a": "clean", "b": ["[NEED METRIC]", {"c": "see <TODO>"}], "d": "CONFIRM this"}
        found = docx_common.scan_placeholders(spec)
        self.assertIn("[NEED METRIC]", found)
        self.assertIn("<TODO>", found)
        self.assertIn("CONFIRM", found)

    def test_clean_content_has_no_stubs(self):
        self.assertEqual(docx_common.scan_placeholders({"a": "all clean here"}), set())

    def test_structural_braces_are_not_stubs(self):
        # only string leaves are scanned — the JSON's own {}/[] are not placeholders
        self.assertEqual(docx_common.scan_placeholders({"x": [1, 2, {"y": "ok"}]}), set())


class TestResumeValidate(unittest.TestCase):
    def _spec(self, title, loc):
        # The bullet is not incidental: a role with none now warns in its own right,
        # so an empty-bulleted fixture could no longer represent a CLEAN spec.
        return {
            "experience": [
                {
                    "company": "Acme",
                    "title": title,
                    "location_dates": loc,
                    "bullets": [["Owned a thing", " that mattered."]],
                }
            ]
        }

    def test_clean_spec_has_no_warnings(self):
        w = resume_builder.validate(
            self._spec("Director of Data Governance", "City, ST | 2020 - 2024")
        )
        self.assertEqual(w, [])

    def test_comma_in_title_warns_about_truncation(self):
        w = resume_builder.validate(
            self._spec("Director, Data Governance", "City, ST | 2020 - 2024")
        )
        self.assertTrue(any("truncate" in x for x in w), w)

    def test_malformed_location_line_warns(self):
        w = resume_builder.validate(self._spec("Director of Data Governance", "Remote"))
        self.assertTrue(any("location/date" in x for x in w), w)

    def test_leftover_placeholder_warns(self):
        spec = self._spec("Director of Data Governance", "City, ST | 2020 - 2024")
        spec["summary"] = "Led [NEED METRIC] teams."
        w = resume_builder.validate(spec)
        self.assertTrue(any("placeholder" in x for x in w), w)


class TestSectionOrderByLevel(unittest.TestCase):
    def _order(self, level):
        return resume_builder.SECTION_ORDERS[level]

    def test_executive_puts_experience_before_tools(self):
        # the record is the pitch at Director/VP; the tool list supports it
        o = self._order("executive")
        self.assertLess(o.index("experience"), o.index("skills"))

    def test_ic_leads_with_tools(self):
        # at IC level the exact-tool match IS the screen
        self.assertEqual(self._order("ic")[0], "skills")

    def test_entry_leads_with_education(self):
        # a thin work history makes the degree the strongest credential
        o = self._order("entry")
        self.assertEqual(o[0], "education")
        self.assertLess(o.index("education"), o.index("experience"))

    def test_education_trails_experience_for_everyone_else(self):
        for level in ("executive", "manager", "ic"):
            o = self._order(level)
            self.assertGreater(
                o.index("education"),
                o.index("experience"),
                f"{level}: education should follow experience",
            )

    def test_entry_targets_one_page_others_two(self):
        self.assertEqual(resume_builder.LEVELS["entry"]["pages"], 1)
        for level in ("executive", "manager", "ic"):
            self.assertEqual(resume_builder.LEVELS[level]["pages"], 2)

    def test_entry_with_competency_grid_warns(self):
        spec = {
            "level": "entry",
            "competencies": ["Leadership | Strategy"],
            "experience": [
                {
                    "company": "Acme",
                    "title": "Analyst",
                    "location_dates": "City, ST | 2024 - 2026",
                    "bullets": [],
                }
            ],
        }
        self.assertTrue(
            any("entry" in w and "Competencies" in w for w in resume_builder.validate(spec))
        )

    def test_every_level_renders_all_four_sections_once(self):
        for level, cfg in resume_builder.LEVELS.items():
            self.assertEqual(
                sorted(cfg["order"]),
                ["competencies", "education", "experience", "skills"],
                f"{level}: must emit each section exactly once",
            )

    def test_competency_grid_stays_above_experience_at_every_level(self):
        # the compact keyword grid is ATS-indexed and orients a human in ~10s
        for level, order in resume_builder.SECTION_ORDERS.items():
            self.assertLess(
                order.index("competencies"),
                order.index("experience"),
                f"{level}: competencies must precede experience",
            )

    def test_default_level_is_executive(self):
        self.assertEqual(resume_builder.DEFAULT_LEVEL, "executive")
        self.assertIn(resume_builder.DEFAULT_LEVEL, resume_builder.SECTION_ORDERS)

    def test_unknown_level_warns(self):
        spec = {
            "level": "principal",  # not a valid key
            "experience": [
                {
                    "company": "Acme",
                    "title": "Director of Data",
                    "location_dates": "City, ST | 2020 - 2024",
                    "bullets": [],
                }
            ],
        }
        self.assertTrue(any("unknown level" in w for w in resume_builder.validate(spec)))

    def test_valid_level_is_silent(self):
        spec = {
            "level": "executive",
            "experience": [
                {
                    "company": "Acme",
                    "title": "Director of Data",
                    "location_dates": "City, ST | 2020 - 2024",
                    "bullets": [],
                }
            ],
        }
        self.assertEqual([w for w in resume_builder.validate(spec) if "level" in w], [])


class TestReverseChronology(unittest.TestCase):
    def _job(self, company, dates):
        return {
            "company": company,
            "title": "Director of Data",
            "location_dates": dates,
            "bullets": [],
        }

    def test_correct_order_is_clean(self):
        spec = {
            "experience": [
                self._job("Northwind Media", "City, ST | April 2024 – Present"),
                self._job("Consulting", "City, ST | September 2023 – April 2026"),
                self._job("Contoso Lending", "City, ST | June 2022 – September 2023"),
                self._job("Hartford", "City, ST | May 2019 – June 2022"),
            ]
        }
        self.assertEqual([w for w in resume_builder.validate(spec) if "chronological" in w], [])

    def test_out_of_order_entry_warns(self):
        # Hartford (2022) placed above Consulting (2026) — the exact bug we hit
        spec = {
            "experience": [
                self._job("Northwind Media", "City, ST | April 2024 – Present"),
                self._job("Hartford", "City, ST | May 2019 – June 2022"),
                self._job("Consulting", "City, ST | September 2023 – April 2026"),
            ]
        }
        w = [x for x in resume_builder.validate(spec) if "reverse-chronological" in x]
        self.assertTrue(w, "expected an out-of-order warning")
        self.assertIn("Consulting", w[0])

    def test_present_sorts_as_most_recent(self):
        # a 'Present' role listed AFTER an ended one is out of order
        spec = {
            "experience": [
                self._job("Old", "City, ST | May 2019 – June 2022"),
                self._job("Current", "City, ST | April 2024 – Present"),
            ]
        }
        self.assertTrue(any("reverse-chronological" in x for x in resume_builder.validate(spec)))

    def test_stacked_subroles_must_descend(self):
        spec = {
            "experience": [
                {
                    "company": "BigCo",
                    "roles": [
                        {
                            "title": "Analyst",
                            "location_dates": "City, ST | 2016 – 2019",
                            "bullets": [],
                        },
                        {
                            "title": "Director",
                            "location_dates": "City, ST | 2019 – 2022",
                            "bullets": [],
                        },
                    ],
                }
            ]
        }
        self.assertTrue(any("stacked roles" in x for x in resume_builder.validate(spec)))

    def test_end_key_parses_formats(self):
        self.assertEqual(resume_builder._end_key("City, ST | May 2019 – June 2022"), (2022, 6))
        self.assertEqual(resume_builder._end_key("City, ST | 2024 – Present"), (9999, 99))
        self.assertEqual(resume_builder._end_key("City, ST | 2020 - 2024"), (2024, 0))


class TestAdvisorySection(unittest.TestCase):
    """Concurrent fractional work goes in its own section, out of the main
    reverse-chron timeline. It was folding into the current job's role-description
    on Workday import because its dates overlapped the current role (2026-07-22)."""

    def _spec(self):
        return {
            "name": "Test Person",
            "contact": "City, ST | test@example.com",
            "level": "executive",
            "summary": "A leader.",
            "competencies": ["A | B"],
            "skills": [["Tools:  ", "X, Y"]],
            "education": [["MS, Field", ", University"]],
            "experience": [
                {
                    "company": "Northwind Media",
                    "title": "Director of Data",
                    "location_dates": "City, ST | April 2024 - Present",
                    "bullets": [["Led", " a team."]],
                },
                {
                    "company": "Contoso Lending",
                    "title": "Head of Data",
                    "location_dates": "City, ST | June 2022 - September 2023",
                    "bullets": [["Built", " a warehouse."]],
                },
            ],
            "advisory": [
                {
                    "company": "Self-Employed",
                    "title": "Fractional Head of Data",
                    "location_dates": "City, ST | September 2023 - April 2026",
                    "bullets": [["Advised", " an EdTech client."]],
                }
            ],
        }

    def test_overlapping_advisory_does_not_trip_reverse_chron(self):
        # the advisory block (ends 2026) sits between the current role and the prior one
        # by date, which WOULD be out of order in the main list. In its own section
        # it is not checked against the main timeline.
        warns = [w for w in resume_builder.validate(self._spec()) if "chronological" in w]
        self.assertEqual(warns, [], warns)

    def test_advisory_entries_are_still_format_checked(self):
        spec = self._spec()
        spec["advisory"][0]["title"] = "Fractional Head, Data"  # comma truncates on import
        self.assertTrue(
            any("truncate" in w for w in resume_builder.validate(spec)),
            "advisory titles must still get the punctuation check",
        )

    def test_advisory_renders_its_own_section(self):
        import os
        import tempfile

        import docx as _docx

        with tempfile.TemporaryDirectory() as d:
            out = os.path.join(d, "r.docx")
            resume_builder.build_resume(self._spec(), out)
            paras = [p.text.strip() for p in _docx.Document(out).paragraphs if p.text.strip()]
        self.assertIn("Advisory & Consulting", paras)
        # the current role is immediately followed by the prior one in the timeline, not by
        # the consulting — that adjacency is the whole point of the move.
        opt, up = paras.index("Northwind Media"), paras.index("Contoso Lending")
        adv = paras.index("Self-Employed")
        self.assertLess(up, adv, "the prior role must precede the advisory block")
        self.assertLess(opt, up)

    def test_no_advisory_key_changes_nothing(self):
        spec = self._spec()
        del spec["advisory"]
        # still valid, still builds — the feature is purely additive
        self.assertEqual([w for w in resume_builder.validate(spec) if "chronological" in w], [])


class TestEarlierFoldedIntoExperience(unittest.TestCase):
    """Older roles render inside the one Professional Experience section, not a
    separate "Earlier Experience" block. Three experience-like sections were a
    suspect in a Workday import dropping the current role (2026-07-22)."""

    def _spec(self):
        return {
            "name": "Test Person",
            "contact": "City, ST | test@example.com",
            "level": "executive",
            "summary": "A leader.",
            "competencies": ["A | B"],
            "skills": [["Tools:  ", "X, Y"]],
            "education": [["MS, Field", ", University"]],
            "experience": [
                {
                    "company": "Northwind Media",
                    "title": "Director of Data",
                    "location_dates": "City, ST | April 2024 - Present",
                    "bullets": [["Led", " a team."]],
                }
            ],
            "earlier": [
                {
                    "company": "Old Employer",
                    "title": "Business Analyst",
                    "location_dates": "City, ST | October 2008 - March 2012",
                }
            ],
        }

    def _paras(self, spec):
        import os
        import tempfile

        import docx as _docx

        with tempfile.TemporaryDirectory() as d:
            out = os.path.join(d, "r.docx")
            resume_builder.build_resume(spec, out)
            return [p.text.strip() for p in _docx.Document(out).paragraphs if p.text.strip()]

    def test_no_earlier_experience_header(self):
        self.assertNotIn("Earlier Experience", self._paras(self._spec()))

    def test_old_role_is_inside_professional_experience(self):
        paras = self._paras(self._spec())
        exp = paras.index("Professional Experience")
        old = paras.index("Old Employer")
        edu = paras.index("Education & Certifications")
        # the old employer sits after the experience header and before Education,
        # with no other section header intervening
        self.assertLess(exp, old)
        self.assertLess(old, edu)

    def test_advisory_still_comes_after_the_folded_experience(self):
        spec = self._spec()
        spec["advisory"] = [
            {
                "company": "Self-Employed",
                "title": "Fractional Head of Data",
                "location_dates": "City, ST | September 2023 - April 2026",
                "bullets": [["Advised", " a client."]],
            }
        ]
        paras = self._paras(spec)
        self.assertLess(paras.index("Old Employer"), paras.index("Advisory & Consulting"))

    def test_no_earlier_key_is_fine(self):
        spec = self._spec()
        del spec["earlier"]
        self.assertNotIn("Earlier Experience", self._paras(spec))


if __name__ == "__main__":
    unittest.main()


class TestEveryRoleCarriesALine(unittest.TestCase):
    """A role listed with only a title and dates gives a reader nothing and an ATS
    nothing to match on — the tenure it was added to prove is all it contributes.

    The validator WARNS rather than filling the gap: the line has to come from the
    profile, and a builder that invented one would be fabricating.
    """

    def _spec(self, **over):
        spec = {
            "name": "A Candidate",
            "contact": "Somewhere, ZZ | a@example.com",
            "level": "executive",
            "experience": [
                {
                    "company": "Acme",
                    "title": "Director of Data",
                    "location_dates": "Somewhere, ZZ | May 2019 – Present",
                    "bullets": [["Did a thing", " that mattered."]],
                }
            ],
        }
        spec.update(over)
        return spec

    def _bare_role_warnings(self, spec):
        return [w for w in resume_builder.validate(spec) if "no content line" in w]

    def test_a_role_with_bullets_is_fine(self):
        self.assertEqual(self._bare_role_warnings(self._spec()), [])

    def test_a_role_with_no_bullets_warns(self):
        spec = self._spec()
        spec["experience"][0]["bullets"] = []
        self.assertEqual(len(self._bare_role_warnings(spec)), 1)

    def test_a_stacked_sub_role_with_no_bullets_warns_and_names_the_employer(self):
        """Sub-roles carry no company of their own — the warning must still say who."""
        spec = self._spec(
            experience=[
                {
                    "company": "Acme",
                    "roles": [
                        {
                            "title": "Director of Data",
                            "location_dates": "Somewhere, ZZ | May 2019 – Present",
                            "bullets": [["Did a thing", " that mattered."]],
                        },
                        {
                            "title": "Data Engineer",
                            "location_dates": "Somewhere, ZZ | June 2013 – May 2019",
                            "bullets": [],
                        },
                    ],
                }
            ]
        )
        warns = self._bare_role_warnings(spec)
        self.assertEqual(len(warns), 1)
        self.assertIn("Acme", warns[0])
        self.assertNotIn("?", warns[0])

    def test_an_earlier_entry_counts_its_singular_bullet(self):
        """`earlier` uses a `bullet` string, not a `bullets` list. Both must count."""
        spec = self._spec(
            earlier=[
                {
                    "company": "Old Co",
                    "title": "Analyst",
                    "location_dates": "Somewhere, ZZ | 2008 – 2012",
                    "bullet": "Automated the reporting set.",
                }
            ]
        )
        self.assertEqual(self._bare_role_warnings(spec), [])

    def test_an_earlier_entry_with_no_bullet_warns(self):
        spec = self._spec(
            earlier=[
                {
                    "company": "Old Co",
                    "title": "Analyst",
                    "location_dates": "Somewhere, ZZ | 2008 – 2012",
                }
            ]
        )
        self.assertEqual(len(self._bare_role_warnings(spec)), 1)

    def test_an_empty_bullet_string_does_not_count(self):
        spec = self._spec(
            earlier=[
                {
                    "company": "Old Co",
                    "title": "Analyst",
                    "location_dates": "Somewhere, ZZ | 2008 – 2012",
                    "bullet": "   ",
                }
            ]
        )
        self.assertEqual(len(self._bare_role_warnings(spec)), 1)

    def test_validation_does_not_mutate_the_spec(self):
        """_all_roles fills company onto COPIES; a synthetic key must never reach
        the spec, where it could be rendered into the document."""
        spec = self._spec(
            experience=[
                {
                    "company": "Acme",
                    "roles": [
                        {
                            "title": "Data Engineer",
                            "location_dates": "Somewhere, ZZ | 2013 – 2019",
                            "bullets": [],
                        }
                    ],
                }
            ]
        )
        resume_builder.validate(spec)
        self.assertNotIn("company", spec["experience"][0]["roles"][0])


class TestACompanyBlockIsNeverSplitAcrossPages(unittest.TestCase):
    """A company header must arrive on the same page as its own content.

    Word breaks between paragraphs wherever the page runs out, so a company name
    could sit alone at the foot of one page with its title, dates and bullets on
    the next. That reads as two different employers and forces the reader back a
    page to work out whose job they are looking at.

    The fix is `keep_with_next` chained down company -> title -> dates, which
    binds the header to at least its first bullet and pushes the whole block over
    instead. It is invisible in the text stream, so nothing else can catch a
    regression: only the paragraph properties say whether it is still on.
    """

    def _doc(self, spec):
        import docx

        out = os.path.join(self.tmp, "r.docx")
        resume_builder.build_resume(spec, out)
        return docx.Document(out)

    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.addCleanup(shutil.rmtree, self.tmp, True)

    def _spec(self, **extra):
        spec = {
            "name": "A Person",
            "contact": "Somewhere, ZZ | a@b.co",
            "summary": "A summary.",
            "experience": [
                {
                    "company": "Acme",
                    "title": "Director of Data",
                    "location_dates": "Somewhere, ZZ | May 2019 – Present",
                    "bullets": [["Did a thing", " and it worked."]],
                }
            ],
        }
        spec.update(extra)
        return spec

    def _bound(self, doc):
        return [p.text for p in doc.paragraphs if p.paragraph_format.keep_with_next]

    def test_company_title_and_dates_are_all_bound_to_what_follows(self):
        bound = self._bound(self._doc(self._spec()))
        self.assertIn("Acme", bound)
        self.assertIn("Director of Data", bound)
        self.assertIn("Somewhere, ZZ | May 2019 – Present", bound)

    def test_a_bullet_is_NOT_bound(self):
        # Only the header chain is bound. Binding bullets too would push whole
        # multi-bullet roles onto a fresh page and waste most of the one before it.
        doc = self._doc(self._spec())
        bullets = [p.text for p in doc.paragraphs if p.text.startswith("•")]
        self.assertTrue(bullets, "expected at least one bullet")
        for b in bullets:
            self.assertNotIn(b, self._bound(doc))

    def test_stacked_sub_roles_each_bind_their_own_header(self):
        spec = self._spec(
            experience=[
                {
                    "company": "Acme",
                    "roles": [
                        {
                            "title": "Director of Data",
                            "location_dates": "Somewhere, ZZ | May 2019 – Present",
                            "bullets": [["Led", " a team."]],
                        },
                        {
                            "title": "Senior Data Engineer",
                            "location_dates": "Somewhere, ZZ | May 2016 – May 2019",
                            "bullets": [["Built", " a pipeline."]],
                        },
                    ],
                }
            ]
        )
        bound = self._bound(self._doc(spec))
        for text in (
            "Acme",
            "Director of Data",
            "Senior Data Engineer",
            "Somewhere, ZZ | May 2016 – May 2019",
        ):
            self.assertIn(text, bound)

    def test_earlier_and_advisory_headers_bind_too(self):
        spec = self._spec(
            earlier=[
                {
                    "company": "Oldco",
                    "title": "Analyst",
                    "location_dates": "Somewhere, ZZ | June 2010 – May 2012",
                    "bullet": "Did early work.",
                }
            ],
            advisory=[
                {
                    "company": "Selfemployed",
                    "title": "Fractional Head of Data",
                    "location_dates": "Somewhere, ZZ | June 2022 – Present",
                    "bullets": [["Advised", " a client."]],
                }
            ],
        )
        bound = self._bound(self._doc(spec))
        for text in ("Oldco", "Analyst", "Selfemployed", "Fractional Head of Data"):
            self.assertIn(text, bound)


class TestBulletContinuationReadsAsOneSentence(unittest.TestCase):
    """The bold lead-in and the rest are ONE sentence, so the continuation opens with
    a space or a comma.

    A colon, semicolon or dash turns the lead-in into a label with an explanation
    hanging off it. Both shapes are defensible in isolation; mixing them inside one
    document is visible to a reader before they have read a word of the content, which
    is what this catches.

    Measured before being enforced: across a real portfolio the sentence shape was
    already overwhelmingly dominant, so the outliers were drift rather than a competing
    house style.
    """

    def _spec(self, rest):
        return {
            "experience": [
                {
                    "company": "Acme",
                    "title": "Director of Data",
                    "location_dates": "Somewhere, ZZ | May 2019 – Present",
                    "bullets": [["Did a thing", rest]],
                }
            ]
        }

    def _warnings(self, rest):
        return [
            w for w in resume_builder.validate(self._spec(rest)) if "continuation starts with" in w
        ]

    def test_space_and_comma_are_the_house_shape(self):
        self.assertEqual(self._warnings(" that mattered."), [])
        self.assertEqual(self._warnings(", which mattered."), [])

    def test_colon_warns(self):
        self.assertEqual(len(self._warnings(": an elaboration follows.")), 1)

    def test_semicolon_warns(self):
        self.assertEqual(len(self._warnings("; an elaboration follows.")), 1)

    def test_spaced_dash_warns(self):
        for dash in (" - ", " – ", " — "):
            self.assertEqual(len(self._warnings(f"{dash}an aside follows.")), 1, dash)

    def test_a_hyphenated_word_is_not_a_separator(self):
        """Only a dash FOLLOWED BY A SPACE separates; 'end-to-end' must not trip it."""
        self.assertEqual(self._warnings(" end-to-end ownership of the platform."), [])

    def test_the_message_names_the_bullet_and_the_fix(self):
        w = self._warnings(": an elaboration follows.")[0]
        self.assertIn("Did a thing", w)
        self.assertIn("Rewrite", w)
